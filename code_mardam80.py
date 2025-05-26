import os
import glob
import pandas as pd
import numpy as np
import math
import torch
import torch.nn as nn
import torch.optim as optim
import torch.nn.functional as F
from torch.distributions import Categorical
import random
from docx import Document

# ======================
# Global Constants & Hyperparameters
# ======================
SERVICE_TIME        = 1.0
TRUCK_SPEED         = 50.0 / 60.0  # distance units per minute
MAX_CAPACITY        = 500
SERVED_BONUS        = 150.0
DEMAND_BONUS        = 2.0
NEW_VEHICLE_PENALTY = 50.0  # to discourage spawning extra vehicles
PENALTY_FACTOR      = 100.0
COMPLETION_BONUS    = 500.0
ENTROPY_COEF        = 0.01
SEED                = 42
GAMMA               = 0.99

# PPO Hyperparams
CLIP_EPS            = 0.2
PPO_EPOCHS          = 4   # fewer epochs
NO_PROGRESS_THRESH  = 50

# ======================
# Seed & Device
# ======================
torch.manual_seed(SEED)
np.random.seed(SEED)
random.seed(SEED)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print("Using device:", device)

# ======================
# Utils: distance, route_distance, 2-opt
# ======================
def euclidean_distance(x1, y1, x2, y2):
    return math.hypot(x1 - x2, y1 - y2)


def route_distance(route, data):
    return sum(euclidean_distance(
        data.iloc[a]['x'], data.iloc[a]['y'],
        data.iloc[b]['x'], data.iloc[b]['y'])
        for a, b in zip(route[:-1], route[1:]))


def two_opt(route, data):
    best = route[:]
    improved = True
    while improved:
        improved = False
        for i in range(1, len(best) - 2):
            for j in range(i + 1, len(best) - 1):
                if j - i == 1: continue
                new_route = best[:]
                new_route[i:j] = reversed(best[i:j])
                if route_distance(new_route, data) < route_distance(best, data):
                    best = new_route
                    improved = True
        route = best
    return best

# ======================
# DSCVRPTW Environment
# ======================
class DSCVRPTWEnv:
    def __init__(self, data, max_vehicles=10):
        self.data = data.reset_index(drop=True)
        self.max_vehicles = max_vehicles
        self.reset()

    def reset(self):
        self.time = 0.0
        self.vehicles = [{
            'route': [0],
            'current_node': 0,
            'capacity': MAX_CAPACITY,
            'time': 0.0,
            'distance': 0.0
        }]
        self.unserved = set(range(1, len(self.data)))
        self.served_flag = np.zeros(len(self.data), dtype=bool)
        self.available = set()
        # initialize wait time trackers
        self.total_vehicle_wait = 0.0
        self.total_customer_wait = 0.0
        return self._get_state()

    def _get_state(self):
        customers = []
        for idx in self.unserved:
            if not self.served_flag[idx]:
                info = self.data.iloc[idx].to_dict()
                info['index'] = idx
                customers.append(info)
        return {'time': self.time, 'vehicles': self.vehicles, 'customers': customers}

    def update_dynamic_customers(self):
        for idx in list(self.unserved):
            if self.time >= self.data.iloc[idx]['time']:
                self.available.add(idx)

    def step(self, actions):
        self.update_dynamic_customers()
        reward = 0.0
        for vidx, cid in actions:
            if cid is None or cid not in self.available or cid not in self.unserved:
                continue
            if self.served_flag[cid]:
                continue
            info = self.data.iloc[cid]
            veh = self.vehicles[vidx]
            if info['demand'] > veh['capacity']:
                continue

            prev = self.data.iloc[veh['current_node']]
            d = euclidean_distance(prev['x'], prev['y'], info['x'], info['y'])
            travel_time = d / TRUCK_SPEED
            # compute vehicle wait if arrives before opening
            depart_time = veh['time'] + travel_time
            vehicle_idle = max(0.0, info['open'] - depart_time)
            arrival = max(depart_time, info['open'])
            # compute customer wait from its appearance to service start
            customer_wait = max(0.0, arrival - info['close'])

            # accumulate wait times
            self.total_vehicle_wait  += vehicle_idle
            self.total_customer_wait += customer_wait

            late = max(0.0, arrival - info['close'])

            # update vehicle state
            veh['route'].append(cid)
            veh['capacity'] -= info['demand']
            veh['time'] = arrival + SERVICE_TIME
            veh['distance'] += d
            veh['current_node'] = cid

            # mark served
            self.unserved.remove(cid)
            self.available.discard(cid)
            self.served_flag[cid] = True

            # reward
            reward += SERVED_BONUS + DEMAND_BONUS * info['demand']
            reward -= d  # distance penalty
            reward -= min(late, PENALTY_FACTOR)

        # advance global time
        next_times = [v['time'] for v in self.vehicles]
        nxt = min(next_times)
        if nxt > self.time:
            self.time = nxt
        else:
            future = [self.data.iloc[i]['time'] for i in self.unserved if self.data.iloc[i]['time'] > self.time]
            if future:
                self.time = min(future)
        self.update_dynamic_customers()

        # spawn new vehicle if none can serve
        if self.unserved:
            can = False
            for veh in self.vehicles:
                cur = self.data.iloc[veh['current_node']]
                for i in self.unserved:
                    if i not in self.available or self.served_flag[i]: continue
                    info = self.data.iloc[i]
                    if info['demand'] > veh['capacity']: continue
                    d = euclidean_distance(cur['x'],cur['y'],info['x'],info['y'])
                    arr = max(veh['time'] + d/TRUCK_SPEED, info['open'])
                    if arr <= info['close']:
                        can = True; break
                if can: break
            if not can and len(self.vehicles) < self.max_vehicles:
                self.vehicles.append({
                    'route': [0], 'current_node': 0,
                    'capacity': MAX_CAPACITY, 'time': self.time, 'distance': 0.0
                })
                reward -= NEW_VEHICLE_PENALTY

        done = not self.unserved
        return self._get_state(), reward, done

    def finalize_routes(self):
        for veh in self.vehicles:
            if veh['current_node'] != 0:
                veh['route'].append(0)
            veh['route'] = two_opt(veh['route'], self.data)
            veh['distance'] = route_distance(veh['route'], self.data)
            veh['current_node'] = 0

    def render(self):
        for i, v in enumerate(self.vehicles):
            print(f"Vehicle {i}: Route={v['route']}, Distance={v['distance']:.2f}, RemCap={v['capacity']}")
        print(f"Total vehicle waiting time   = {self.total_vehicle_wait:.2f}")
        print(f"Total customer waiting time  = {self.total_customer_wait:.2f}")

    def total_distance(self):
        return sum(v['distance'] for v in self.vehicles)

    def get_wait_times(self):
        return self.total_vehicle_wait, self.total_customer_wait

# ======================
# MARDAM & ActorCritic
# ======================
class MARDAM(nn.Module):
    def __init__(self, cdim, vdim, embed=128, heads=8, layers=2):
        super().__init__()
        self.ce = nn.Linear(cdim, embed)
        self.ve = nn.Linear(vdim, embed)
        enc = nn.TransformerEncoderLayer(d_model=embed, nhead=heads, batch_first=True)
        self.tr = nn.TransformerEncoder(enc, num_layers=layers)
        self.scorer = nn.Linear(embed*2, 1)
    def forward(self, cf, vf):
        ce = self.tr(self.ce(cf).unsqueeze(0)).squeeze(0)
        ve = self.ve(vf)
        M,N = ve.size(0), ce.size(0)
        ve_e = ve.unsqueeze(1).expand(M,N,-1)
        ce_e = ce.unsqueeze(0).expand(M,N,-1)
        scores = self.scorer(torch.cat([ve_e,ce_e],-1)).squeeze(-1)
        return F.softmax(scores, dim=-1)

class ActorCritic(nn.Module):
    def __init__(self, cdim, vdim, embed=128, heads=8, layers=2):
        super().__init__()
        self.actor = MARDAM(cdim, vdim, embed, heads, layers)
        self.ce    = nn.Linear(cdim, embed)
        enc       = nn.TransformerEncoderLayer(d_model=embed, nhead=heads, batch_first=True)
        self.critic_tr = nn.TransformerEncoder(enc, num_layers=layers)
        self.pool = nn.AdaptiveAvgPool1d(1)
        self.head = nn.Linear(embed,1)
    def forward(self, cf, vf):
        p = self.actor(cf, vf)
        x = self.critic_tr(self.ce(cf).unsqueeze(0)).transpose(1,2)
        v = self.head(self.pool(x).squeeze(-1)).squeeze(0)
        return p, v

# ======================
# Feature Extraction
# ======================
def extract_features(state, data):
    vf, cf = [], []
    custs = state['customers']
    # vehicle features
    for v in state['vehicles']:
        n = data.iloc[v['current_node']]
        vf.append([n['x'],n['y'],v['capacity'],v['time']])
    # static: time==0 and available
    static_idxs = [i for i,c in enumerate(custs) if c['time']==0]
    # dynamic: time>0 and current time>=c['time']
    dynamic_idxs = [i for i,c in enumerate(custs) if c['time']>0 and state['time']>=c['time']]
    if static_idxs:
        mask = [i in static_idxs for i in range(len(custs))]
    else:
        mask = [i in dynamic_idxs for i in range(len(custs))]
    # customer features
    for c in custs:
        cf.append([c['x'],c['y'],c['demand'],c['open'],c['close'],c['time']])
    cf = torch.tensor(cf, dtype=torch.float)
    vf = torch.tensor(vf, dtype=torch.float)
    mask = torch.tensor(mask, dtype=torch.bool)
    return cf.to(device), vf.to(device), mask.to(device), custs

# ======================
# PPO Update
# ======================
def ppo_update(transitions, model, opt):
    advs = torch.stack([t['advantage'] for t in transitions]).to(device)
    advs = (advs - advs.mean())/(advs.std()+1e-8)
    total_loss=0.0
    for _ in range(PPO_EPOCHS):
        for t in transitions:
            cf,vf,mask = t['cust_feats'].to(device), t['veh_feats'].to(device), t['mask'].to(device)
            old_lp = t['log_prob'].to(device)
            ret    = t['return'].to(device)
            adv    = t['advantage'].to(device)
            probs,val = model(cf,vf)
            probs = probs*mask.float()
            probs = probs/(probs.sum(-1,keepdim=True)+1e-10)
            lps,ents=[],[]
            for vidx,a in enumerate(t['actions']):
                if a is not None:
                    dist = Categorical(probs[vidx])
                    lps.append(dist.log_prob(torch.tensor(a,device=device)))
                    ents.append(dist.entropy())
            if not lps: continue
            new_lp = torch.stack(lps).mean()
            ent = torch.stack(ents).mean()
            ratio = torch.exp(new_lp-old_lp)
            s1 = ratio*adv; s2 = torch.clamp(ratio,1-CLIP_EPS,1+CLIP_EPS)*adv
            a_loss = -torch.min(s1,s2)
            c_loss = F.mse_loss(val, ret.unsqueeze(0))
            loss   = a_loss + c_loss + ENTROPY_COEF*(-ent)
            opt.zero_grad(); loss.backward(); torch.nn.utils.clip_grad_norm_(model.parameters(),0.5); opt.step()
            total_loss+=loss.item()
    return total_loss

# ======================
# Training Loop
# ======================
def train(model, env, data, n_eps=50, lr=1e-4, wdecay=1e-5, max_steps=1000):
    opt = optim.AdamW(model.parameters(), lr=lr, weight_decay=wdecay)
    sched = optim.lr_scheduler.ReduceLROnPlateau(opt, 'max', factor=0.5, patience=10)
    for ep in range(1,n_eps+1):
        state=env.reset(); transitions=[]; steps=0; no_p=0; last_s=0; tot_r=0.0
        while steps<max_steps:
            steps+=1
            cf,vf,mask,custs = extract_features(state,data)
            if cf.size(0)==0: break
            probs,val = model(cf,vf)
            probs = probs*mask.float(); probs/=probs.sum(-1,keepdim=True)
            actions,lps=[],[]
            for vidx in range(probs.size(0)):
                # nếu không có khách khả dụng nào thì xe đứng yên
                if mask.sum() == 0:
                    actions.append(None)
                    lps.append(torch.tensor(0.0, device=device))
                else:
                    dist = Categorical(probs[vidx])
                    a    = dist.sample()
                    actions.append(a.item())
                    lps.append(dist.log_prob(a))

            env_act=[(i,custs[a]['index'] if a is not None else None) for i,a in enumerate(actions)]
            nxt,rew,done=env.step(env_act); tot_r+=rew
            transitions.append({
                'cust_feats':cf.detach().cpu(), 'veh_feats':vf.detach().cpu(), 'mask':mask.detach().cpu(),
                'actions':actions, 'log_prob':torch.stack(lps).mean().detach(),
                'reward':rew, 'state_value':val.detach() })
            state=nxt
            cur_s = int(env.served_flag.sum())
            no_p = no_p+1 if cur_s==last_s else 0; last_s=cur_s
            if no_p>=NO_PROGRESS_THRESH: break
        tot_r += COMPLETION_BONUS if not env.unserved else -PENALTY_FACTOR*len(env.unserved)
        R=0.0
        for t in reversed(transitions): R=t['reward']+GAMMA*R; t['return']=torch.tensor(R,dtype=torch.float)
        for t in transitions: t['advantage']=t['return']-t['state_value']
        loss=ppo_update(transitions,model,opt)
        env.finalize_routes(); dist=env.total_distance()
        print(f"Ep {ep}/{n_eps} | Reward {tot_r:.2f} | Loss {loss:.2f} | Dist {dist:.2f} | Vehl {len(env.vehicles)}")
        env.render(); print('-'*50)
        sched.step(tot_r); torch.cuda.empty_cache()
    return model

if __name__ == "__main__":
    # Thư mục chứa CSV đầu vào và thư mục xuất DOCX
    input_dir  = "F:/KLTN/80_Customer"
    output_dir = "F:/KLTN/code/output80"
    os.makedirs(output_dir, exist_ok=True)

    # Duyệt tất cả .csv trong input_dir
    csv_files = glob.glob(os.path.join(input_dir, "*.csv"))
    if not csv_files:
        print(f"No CSV files found in {input_dir}")
        exit()

    for csv_path in csv_files:
        filename = os.path.splitext(os.path.basename(csv_path))[0]
        print(f"\n=== Processing {filename}.csv ===")
        data = pd.read_csv(csv_path)

        # Tạo document mới
        doc = Document()
        doc.add_heading(f"Results for {filename}.csv", level=1)

        results = {}
        
        for max_v in range(5, 9):
            doc.add_heading(f"max_vehicles = {max_v}", level=2)
            print(f"\n--- Training with max_vehicles = {max_v} ---")
            env   = DSCVRPTWEnv(data, max_vehicles=max_v)
            model = ActorCritic(cdim=6, vdim=4).to(device)

            trained_model = train(model, env, data)

            # Sau khi train xong, finalize và lấy số liệu
            env.finalize_routes()
            total_dist = env.total_distance()
            v_wait, c_wait = env.get_wait_times()
            results[max_v] = total_dist

            # Ghi vào document
            p = doc.add_paragraph()
            p.add_run(f"Total distance = {total_dist:.2f}\n").bold = True
            p.add_run(f"Vehicle total idle time   = {v_wait:.2f}\n")
            p.add_run(f"Customer total wait time = {c_wait:.2f}\n")

        # Summary cuối
        doc.add_heading("Summary of total distances", level=2)
        for max_v, dist in results.items():
            doc.add_paragraph(f"• max_vehicles = {max_v} → total_distance = {dist:.2f}")

        # Lưu file docx
        output_path = os.path.join(output_dir, f"{filename}_results.docx")
        doc.save(output_path)
        print(f"Saved results to {output_path}")